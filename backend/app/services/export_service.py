from pathlib import Path
from docx import Document
from ..models import Novel, Chapter


class ExportService:
    def __init__(self, export_dir: str):
        self.export_dir = Path(export_dir)
        self.export_dir.mkdir(parents=True, exist_ok=True)

    def export_txt(self, novel: Novel, chapters: list[Chapter]) -> Path:
        file_path = self.export_dir / f"{novel.id}_{novel.title}_投稿稿件.txt"
        content = [novel.title, "", f"简介：{novel.synopsis}", f"标签：{novel.tags}", ""]
        for chapter in sorted(chapters, key=lambda item: item.chapter_number):
            content.append(chapter.title)
            content.append(chapter.content or "")
            content.append("")
        file_path.write_text("\n".join(content), encoding="utf-8")
        return file_path

    def export_docx(self, novel: Novel, chapters: list[Chapter]) -> Path:
        file_path = self.export_dir / f"{novel.id}_{novel.title}_投稿稿件.docx"
        doc = Document()
        doc.add_heading(novel.title, 0)
        doc.add_paragraph(f"简介：{novel.synopsis}")
        doc.add_paragraph(f"标签：{novel.tags}")
        for chapter in sorted(chapters, key=lambda item: item.chapter_number):
            doc.add_heading(chapter.title, level=1)
            for paragraph in (chapter.content or "").split("\n"):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
        doc.save(file_path)
        return file_path
